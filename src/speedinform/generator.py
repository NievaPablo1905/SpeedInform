"""
Generador de documentos DOCX desde plantilla.

Flujo:
1. Carga la plantilla DOCX.
2. Reemplaza marcadores {{campo}} en todos los párrafos y tablas.
   Para manejar marcadores divididos entre runs, une el texto del párrafo,
   detecta el marcador y reconstruye el primer run con el texto reemplazado.
3. Divide el documento por marcadores de sección:
   [[INICIO_INFORME]], [[INICIO_AVOCAMIENTO]], [[INICIO_ELEVACION]]
4. Guarda tres DOCX: informe.docx, avocamiento.docx, elevacion.docx.
"""

import os
import copy
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.oxml.ns import qn
    import lxml.etree as etree
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.error("python-docx no está instalado.")


SECTION_MARKERS = [
    "[[INICIO_INFORME]]",
    "[[INICIO_AVOCAMIENTO]]",
    "[[INICIO_ELEVACION]]",
]

SECTION_NAMES = {
    "[[INICIO_INFORME]]": "informe.docx",
    "[[INICIO_AVOCAMIENTO]]": "avocamiento.docx",
    "[[INICIO_ELEVACION]]": "elevacion.docx",
}


def _find_template(template_path: str | None = None) -> str:
    """
    Localiza la plantilla DOCX.
    Busca en: argumento provisto → cwd → raíz del paquete → directorio padre del paquete.
    """
    candidates = []
    if template_path:
        candidates.append(template_path)

    pkg_dir = Path(__file__).parent
    repo_root = pkg_dir.parent.parent  # src/speedinform → src → repo root

    template_name = "PLANTILLA INFORME PARA PROGRAM.docx"
    candidates += [
        Path(os.getcwd()) / template_name,
        repo_root / template_name,
        pkg_dir / template_name,
    ]

    for c in candidates:
        if Path(c).exists():
            return str(c)

    raise FileNotFoundError(
        f"No se encontró la plantilla '{template_name}'. "
        f"Candidatos revisados: {candidates}"
    )


def _paragraph_full_text(paragraph) -> str:
    """Devuelve el texto completo de un párrafo."""
    return "".join(run.text for run in paragraph.runs)


def _replace_in_paragraph(paragraph, fields: dict) -> None:
    """
    Reemplaza marcadores {{key}} en el párrafo.
    Une los runs para detectar marcadores que estén divididos,
    luego reconstruye el primer run con el texto completo y limpia los demás.
    """
    full_text = _paragraph_full_text(paragraph)
    replaced_text = full_text
    for key, value in fields.items():
        placeholder = f"{{{{{key}}}}}"
        replaced_text = replaced_text.replace(placeholder, str(value) if value else "")

    if replaced_text == full_text:
        return  # Nada que cambiar

    if paragraph.runs:
        paragraph.runs[0].text = replaced_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        # Párrafo sin runs: agregar uno
        paragraph.add_run(replaced_text)


def _replace_in_document(doc, fields: dict) -> None:
    """Reemplaza marcadores en todos los párrafos y tablas del documento."""
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, fields)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, fields)


def _paragraph_marker(paragraph) -> str | None:
    """Devuelve el marcador de sección si el párrafo lo contiene, o None."""
    text = _paragraph_full_text(paragraph).strip()
    for marker in SECTION_MARKERS:
        if marker in text:
            return marker
    return None


def _copy_paragraph_to_doc(src_paragraph, dst_doc) -> None:
    """Copia un párrafo al documento destino preservando formato XML."""
    new_para = copy.deepcopy(src_paragraph._element)
    dst_doc.element.body.append(new_para)


def _split_document(doc, template_path: str, output_dir: str) -> list[str]:
    """
    Divide el documento por marcadores de sección y guarda cada sección como DOCX.
    """
    # Recolectar todos los párrafos del body (incluyendo los de tablas no, sólo los directos)
    paragraphs = doc.paragraphs
    body_elements = list(doc.element.body)

    # Mapear párrafos a sus posiciones en el body
    # Necesitamos iterar sobre los elementos directos del body
    current_marker = None
    sections: dict[str, list] = {m: [] for m in SECTION_MARKERS}

    for elem in body_elements:
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "p":
            # Es un párrafo
            text = "".join(
                r.text for r in elem.findall(f".//{qn('w:t')}") if r.text
            ).strip()
            marker_found = None
            for marker in SECTION_MARKERS:
                if marker in text:
                    marker_found = marker
                    break
            if marker_found:
                current_marker = marker_found
                continue  # No incluir el marcador en el output
        elif tag == "tbl":
            # Es una tabla
            pass

        if current_marker:
            sections[current_marker].append(copy.deepcopy(elem))

    output_paths = []
    for marker, filename in SECTION_NAMES.items():
        elements = sections[marker]
        if not elements:
            logger.warning("Sección '%s' está vacía; archivo no generado.", marker)
            continue

        # Crear nuevo documento basado en la plantilla (preserva estilos)
        new_doc = Document(template_path)
        # Limpiar body existente
        body = new_doc.element.body
        for child in list(body):
            body.remove(child)

        for elem in elements:
            body.append(elem)

        out_path = os.path.join(output_dir, filename)
        new_doc.save(out_path)
        output_paths.append(out_path)
        logger.info("Documento guardado: %s", out_path)

    return output_paths


def generate_documents(
    template_path: str | None,
    fields: dict,
    output_dir: str,
) -> list[str]:
    """
    Genera los tres documentos DOCX desde la plantilla.

    Args:
        template_path: Ruta a la plantilla DOCX (None para búsqueda automática).
        fields: Diccionario con los campos a sustituir.
        output_dir: Directorio donde se guardarán los archivos generados.

    Returns:
        Lista de rutas a los archivos DOCX generados.
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx no está instalado.")

    template_path = _find_template(template_path)
    os.makedirs(output_dir, exist_ok=True)

    doc = Document(template_path)
    _replace_in_document(doc, fields)

    return _split_document(doc, template_path, output_dir)
